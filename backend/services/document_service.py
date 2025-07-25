"""
Document processing service for RFP documents.
Handles file upload, local storage, and text extraction using Azure Document Intelligence.
"""
import os
import uuid
import aiofiles
from pathlib import Path
from typing import List, Optional
from datetime import datetime

from azure.ai.documentintelligence.aio import DocumentIntelligenceClient
from azure.ai.documentintelligence.models import DocumentAnalysisFeature, AnalyzeResult
from azure.core.credentials import AzureKeyCredential

from models import Document, DocumentCreate
from database_config import database, get_table_name
from services.file_storage_service import file_storage_service


class DocumentService:
    """Service for document upload, storage, and processing with local file storage."""
    
    def __init__(self):
        # Azure Document Intelligence setup
        self.doc_intelligence_endpoint = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_ENDPOINT")
        self.doc_intelligence_key = os.getenv("AZURE_DOCUMENT_INTELLIGENCE_KEY")
        
        if self.doc_intelligence_endpoint and self.doc_intelligence_key:
            self.doc_client = DocumentIntelligenceClient(
                endpoint=self.doc_intelligence_endpoint,
                credential=AzureKeyCredential(self.doc_intelligence_key)
            )
        else:
            self.doc_client = None
            print("Warning: Azure Document Intelligence not configured")
        
        # File storage service
        self.file_storage = file_storage_service
        
        # Allowed file types
        self.allowed_extensions = {'.pdf', '.docx', '.doc', '.txt', '.xlsx', '.pptx'}
        self.max_file_size = 100 * 1024 * 1024  # 100MB
    
    async def upload_document(
        self, 
        file_content: bytes, 
        filename: str, 
        project_id: str,
        content_type: str = None
    ) -> Document:
        """Upload and process a document using local file storage."""
        
        # Get organization_id from project (needed for file storage)
        # For now, we'll use the project_id as organization_id
        # In a real implementation, you'd query the database to get the organization_id
        organization_id = project_id  # Temporary - should be fetched from database
        
        # Generate document ID
        document_id = str(uuid.uuid4())
        
        # Store file using local file storage service
        storage_result = await self.file_storage.store_document(
            file_content=file_content,
            filename=filename,
            organization_id=organization_id,
            project_id=project_id,
            document_id=document_id
        )
        
        if not storage_result["success"]:
            raise ValueError("Failed to store document")
        
        # Create document record
        document = Document(
            id=uuid.UUID(document_id),
            name=filename,
            original_name=filename,
            file_path=storage_result["relative_path"],
            file_size=storage_result["metadata"]["file_size"],
            file_type=storage_result["metadata"]["mime_type"],
            project_id=uuid.UUID(project_id),
            uploaded_at=datetime.now(),
            status="uploaded"
        )
        
        # Save to database
        await self._save_document(document)
        
        # Process document asynchronously
        try:
            processed_content = await self._process_document(storage_result["file_path"], document.file_type)
            document.status = "processed"
            document.processed_at = datetime.now()
            await self._update_document_status(document.id, "processed", processed_content)
        except Exception as e:
            print(f"Document processing failed: {e}")
            document.status = "error"
            await self._update_document_status(document.id, "error")
        
        return document
    
    async def _save_document(self, document: Document):
        """Save document to database."""
        doc_table = get_table_name("documents")
        query = f"""
            INSERT INTO {doc_table} 
            (id, name, original_name, file_path, file_size, file_type, project_id, uploaded_at, status)
            VALUES (:id, :name, :original_name, :file_path, :file_size, :file_type, :project_id, :uploaded_at, :status)
        """
        await database.execute(query, {
            "id": str(document.id),
            "name": document.name,
            "original_name": document.original_name,
            "file_path": document.file_path,
            "file_size": document.file_size,
            "file_type": document.file_type,
            "project_id": str(document.project_id),
            "uploaded_at": document.uploaded_at,
            "status": document.status
        })
    
    async def _update_document_status(self, doc_id: uuid.UUID, status: str, content: str = None):
        """Update document processing status."""
        doc_table = get_table_name("documents")
        
        if content:
            # Store processed content (you might want to add a content field to documents table)
            query = f"""
                UPDATE {doc_table} 
                SET status = :status, processed_at = :processed_at 
                WHERE id = :doc_id
            """
            await database.execute(query, {
                "status": status,
                "processed_at": datetime.now(),
                "doc_id": str(doc_id)
            })
        else:
            query = f"""
                UPDATE {doc_table} 
                SET status = :status 
                WHERE id = :doc_id
            """
            await database.execute(query, {
                "status": status,
                "doc_id": str(doc_id)
            })
    
    async def _process_document(self, file_path: Path, file_type: str) -> str:
        """Process document to extract text content."""
        if not self.doc_client:
            # Fallback: simple text extraction
            return await self._simple_text_extraction(file_path, file_type)
        
        try:
            # Use Azure Document Intelligence for advanced processing
            with open(file_path, 'rb') as f:
                file_content = f.read()
            
            # Analyze document
            poller = await self.doc_client.begin_analyze_document(
                "prebuilt-layout",
                file_content,
                features=[DocumentAnalysisFeature.OCR]
            )
            result = await poller.result()
            
            # Extract text content
            content = ""
            if result.content:
                content = result.content
            
            return content
            
        except Exception as e:
            print(f"Azure Document Intelligence processing failed: {e}")
            # Fallback to simple extraction
            return await self._simple_text_extraction(file_path, file_type)
    
    async def _simple_text_extraction(self, file_path: Path, file_type: str) -> str:
        """Fallback text extraction for when Azure DI is not available."""
        if file_type == '.txt':
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
                return await f.read()
        elif file_type == '.pdf':
            try:
                import PyPDF2
                text = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                return text
            except Exception as e:
                print(f"PDF extraction failed: {e}")
                return f"[PDF processing failed: {str(e)}]"
        elif file_type in ['.docx', '.doc']:
            try:
                import docx
                doc = docx.Document(file_path)
                text = ""
                for paragraph in doc.paragraphs:
                    text += paragraph.text + "\n"
                return text
            except Exception as e:
                print(f"DOCX extraction failed: {e}")
                return f"[DOCX processing failed: {str(e)}]"
        elif file_type in ['.xlsx', '.xls']:
            try:
                import openpyxl
                text = ""
                workbook = openpyxl.load_workbook(file_path, data_only=True)
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"--- Sheet: {sheet_name} ---\n"
                    for row in sheet.iter_rows(values_only=True):
                        row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip():  # Only add non-empty rows
                            text += row_text + "\n"
                    text += "\n"
                return text
            except Exception as e:
                print(f"Excel extraction failed: {e}")
                return f"[Excel processing failed: {str(e)}]"
        else:
            return f"[Unsupported file type for text extraction: {file_type}]"
    
    async def get_project_documents(self, project_id: str) -> List[Document]:
        """Get all documents for a project."""
        doc_table = get_table_name("documents")
        query = f"""
            SELECT * FROM {doc_table} 
            WHERE project_id = :project_id 
            ORDER BY uploaded_at DESC
        """
        rows = await database.fetch_all(query, {"project_id": project_id})
        return [Document(**dict(row)) for row in rows]
    
    async def get_document(self, doc_id: str) -> Optional[Document]:
        """Get a specific document."""
        doc_table = get_table_name("documents")
        query = f"SELECT * FROM {doc_table} WHERE id = :doc_id"
        row = await database.fetch_one(query, {"doc_id": doc_id})
        
        if not row:
            return None
        
        return Document(**dict(row))
    
    async def delete_document(self, doc_id: str) -> bool:
        """Delete a document and its file."""
        document = await self.get_document(doc_id)
        if not document:
            return False
        
        # Delete file from disk
        try:
            file_path = Path(document.file_path)
            if file_path.exists():
                file_path.unlink()
        except Exception as e:
            print(f"Failed to delete file {document.file_path}: {e}")
        
        # Delete from database
        doc_table = get_table_name("documents")
        await database.execute(
            f"DELETE FROM {doc_table} WHERE id = :doc_id",
            {"doc_id": doc_id}
        )
        
        return True


# Global service instance
document_service = DocumentService()